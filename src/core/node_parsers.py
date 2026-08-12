"""
Format parsers that convert various data formats into Nodes.

Each parser:
1. Detects if it can handle the input
2. Parses input into a list of Nodes
3. Can reconstruct from Nodes back to the format
"""

from typing import Any, List, Optional, Union
from abc import ABC, abstractmethod
import json
import re
import io

from .node import Node, NodeType, ParseResult


class BaseNodeParser(ABC):
    """Base class for all node parsers"""

    @property
    @abstractmethod
    def format_name(self) -> str:
        """Name of the format (e.g., 'json', 'markdown')"""
        pass

    @property
    @abstractmethod
    def priority(self) -> int:
        """Priority (lower = try first)"""
        pass

    @abstractmethod
    def can_parse(self, data: Any, format_hint: Optional[str] = None) -> bool:
        """Check if this parser can handle the data"""
        pass

    @abstractmethod
    def parse(self, data: Any) -> ParseResult:
        """Parse data into nodes"""
        pass

    @abstractmethod
    def reconstruct(self, nodes: List[Node], target_format: Optional[str] = None) -> Any:
        """Reconstruct data from nodes"""
        pass


# ============================================================================
# JSON Parser
# ============================================================================

class JSONNodeParser(BaseNodeParser):
    """Parse JSON/dict data into nodes"""

    @property
    def format_name(self) -> str:
        return "json"

    @property
    def priority(self) -> int:
        return 10

    def can_parse(self, data: Any, format_hint: Optional[str] = None) -> bool:
        if format_hint == "json":
            return True
        if isinstance(data, dict):
            return True
        if isinstance(data, str):
            try:
                json.loads(data)
                return True
            except (json.JSONDecodeError, TypeError):
                return False
        return False

    def parse(self, data: Any) -> ParseResult:
        """Parse JSON into nodes"""
        try:
            # Parse string to dict if needed
            if isinstance(data, str):
                data = json.loads(data)

            nodes = self._extract_nodes(data, "")
            return ParseResult(
                nodes=nodes,
                format_name="json",
                success=True,
                metadata={"node_count": len(nodes)}
            )
        except Exception as e:
            return ParseResult(
                nodes=[],
                format_name="json",
                success=False,
                error=str(e)
            )

    def _extract_nodes(self, data: Any, path: str) -> List[Node]:
        """Recursively extract nodes from JSON structure"""
        nodes = []

        if isinstance(data, dict):
            # Check if this is a leaf object (no nested dicts/arrays with dicts)
            has_complex_children = any(
                isinstance(v, dict) or (isinstance(v, list) and v and isinstance(v[0], dict))
                for v in data.values()
            )

            if not has_complex_children:
                # Leaf object - create single node
                nodes.append(Node(
                    path=path or "root",
                    content=data,
                    node_type=NodeType.OBJECT,
                    metadata={"format": "json"}
                ))
            else:
                # Has complex children - recurse
                for key, value in data.items():
                    child_path = f"{path}.{key}" if path else key

                    if isinstance(value, list) and value and isinstance(value[0], dict):
                        # Array of objects - extract each
                        for idx, item in enumerate(value):
                            nodes.extend(self._extract_nodes(item, f"{child_path}[{idx}]"))
                    elif isinstance(value, dict):
                        # Nested object - recurse
                        nodes.extend(self._extract_nodes(value, child_path))
                    # else: primitive values stay with parent object

        elif isinstance(data, list):
            # Top-level array
            for idx, item in enumerate(data):
                item_path = f"{path}[{idx}]" if path else f"[{idx}]"
                nodes.extend(self._extract_nodes(item, item_path))

        else:
            # Primitive value
            nodes.append(Node(
                path=path or "value",
                content=data,
                node_type=NodeType.PRIMITIVE,
                metadata={"format": "json"}
            ))

        return nodes

    def reconstruct(self, nodes: List[Node], target_format: Optional[str] = None) -> Any:
        """Reconstruct JSON from nodes"""
        if not nodes:
            return {}

        # Build nested structure from paths
        result = {}

        for node in nodes:
            self._set_by_path(result, node.path, node.content)

        # If target format is string, dump as JSON
        if target_format == "string":
            return json.dumps(result, indent=2, ensure_ascii=False)

        return result

    def _set_by_path(self, obj: dict, path: str, value: Any):
        """Set value in nested dict by path"""
        if not path or path == "root":
            if isinstance(value, dict):
                obj.update(value)
            return

        parts = []
        current = ""
        in_bracket = False

        # Parse path like "a.b[0].c"
        for char in path:
            if char == '[':
                if current:
                    parts.append(current)
                    current = ""
                in_bracket = True
            elif char == ']':
                if current:
                    parts.append(f"[{current}]")
                    current = ""
                in_bracket = False
            elif char == '.' and not in_bracket:
                if current:
                    parts.append(current)
                    current = ""
            else:
                current += char

        if current:
            parts.append(current)

        # Navigate/create structure
        current_obj = obj
        for i, part in enumerate(parts[:-1]):
            if part.startswith('['):
                # Array index
                idx = int(part[1:-1])
                if not isinstance(current_obj, list):
                    current_obj = []
                while len(current_obj) <= idx:
                    current_obj.append({})
                current_obj = current_obj[idx]
            else:
                # Object key
                if part not in current_obj:
                    # Look ahead to see if next is array
                    if i + 1 < len(parts) and parts[i + 1].startswith('['):
                        current_obj[part] = []
                    else:
                        current_obj[part] = {}
                current_obj = current_obj[part]

        # Set final value
        final_part = parts[-1] if parts else None
        if final_part:
            if final_part.startswith('['):
                idx = int(final_part[1:-1])
                if not isinstance(current_obj, list):
                    current_obj = []
                while len(current_obj) <= idx:
                    current_obj.append(None)
                current_obj[idx] = value
            else:
                if isinstance(value, dict):
                    if final_part not in current_obj:
                        current_obj[final_part] = {}
                    current_obj[final_part].update(value)
                else:
                    current_obj[final_part] = value


# ============================================================================
# YAML Parser
# ============================================================================

class YAMLNodeParser(JSONNodeParser):
    """Parse YAML data (reuses JSON logic since YAML → dict)"""

    @property
    def format_name(self) -> str:
        return "yaml"

    @property
    def priority(self) -> int:
        return 11

    def can_parse(self, data: Any, format_hint: Optional[str] = None) -> bool:
        if format_hint == "yaml":
            return True
        if isinstance(data, str) and not data.strip().startswith('{'):
            # YAML has specific patterns: key: value, lists with -, etc.
            # Don't catch CSV (has commas) or plain text (no colons)
            if ',' in data and '\n' in data:
                # Likely CSV, not YAML
                return False
            if ':' not in data and '-' not in data:
                # No YAML markers, probably plain text
                return False
            try:
                import yaml
                result = yaml.safe_load(data)
                # Only accept if it parses to dict or list (not plain string)
                return isinstance(result, (dict, list))
            except:
                return False
        return False

    def parse(self, data: Any) -> ParseResult:
        """Parse YAML into nodes"""
        try:
            if isinstance(data, str):
                import yaml
                data = yaml.safe_load(data)

            nodes = self._extract_nodes(data, "")
            return ParseResult(
                nodes=nodes,
                format_name="yaml",
                success=True,
                metadata={"node_count": len(nodes)}
            )
        except Exception as e:
            return ParseResult(
                nodes=[],
                format_name="yaml",
                success=False,
                error=str(e)
            )

    def reconstruct(self, nodes: List[Node], target_format: Optional[str] = None) -> Any:
        """Reconstruct as YAML string"""
        result = super().reconstruct(nodes, target_format=None)

        if target_format == "string":
            import yaml
            return yaml.dump(result, default_flow_style=False, allow_unicode=True)

        return result


# ============================================================================
# Plain Text Parser
# ============================================================================

class PlainTextNodeParser(BaseNodeParser):
    """Parse plain text into sentence/paragraph nodes"""

    @property
    def format_name(self) -> str:
        return "text"

    @property
    def priority(self) -> int:
        return 100  # Fallback parser

    def can_parse(self, data: Any, format_hint: Optional[str] = None) -> bool:
        return isinstance(data, str)

    def parse(self, data: Any) -> ParseResult:
        """Parse text into nodes"""
        try:
            text = str(data).strip()

            if not text:
                return ParseResult(
                    nodes=[],
                    format_name="text",
                    success=True
                )

            nodes = []

            # Split into paragraphs
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

            if len(paragraphs) == 1:
                # Single paragraph - split into sentences
                sentences = self._split_sentences(paragraphs[0])
                for i, sentence in enumerate(sentences):
                    nodes.append(Node(
                        path=f"sentence_{i}",
                        content=sentence,
                        node_type=NodeType.SENTENCE,
                        metadata={"format": "text"}
                    ))
            else:
                # Multiple paragraphs
                for i, para in enumerate(paragraphs):
                    nodes.append(Node(
                        path=f"paragraph_{i}",
                        content=para,
                        node_type=NodeType.PARAGRAPH,
                        metadata={"format": "text"}
                    ))

            return ParseResult(
                nodes=nodes,
                format_name="text",
                success=True,
                metadata={"node_count": len(nodes)}
            )
        except Exception as e:
            return ParseResult(
                nodes=[],
                format_name="text",
                success=False,
                error=str(e)
            )

    def _split_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        # Simple sentence splitting
        sentences = re.split(r'[.!?]+\s+', text)
        return [s.strip() for s in sentences if s.strip()]

    def reconstruct(self, nodes: List[Node], target_format: Optional[str] = None) -> str:
        """Reconstruct text from nodes"""
        if not nodes:
            return ""

        # Group by paragraph if available
        paragraphs = {}
        for node in nodes:
            if node.node_type == NodeType.PARAGRAPH:
                # Extract paragraph number
                match = re.match(r'paragraph_(\d+)', node.path)
                if match:
                    idx = int(match.group(1))
                    paragraphs[idx] = str(node.content)
            elif node.node_type == NodeType.SENTENCE:
                # Group sentences into paragraph 0
                if 0 not in paragraphs:
                    paragraphs[0] = []
                if not isinstance(paragraphs[0], list):
                    paragraphs[0] = [paragraphs[0]]
                paragraphs[0].append(str(node.content))

        # Reconstruct
        result = []
        for idx in sorted(paragraphs.keys()):
            content = paragraphs[idx]
            if isinstance(content, list):
                result.append(' '.join(content))
            else:
                result.append(content)

        return '\n\n'.join(result)


# ============================================================================
# Markdown Parser
# ============================================================================

class MarkdownNodeParser(BaseNodeParser):
    """Parse Markdown into structured nodes"""

    @property
    def format_name(self) -> str:
        return "markdown"

    @property
    def priority(self) -> int:
        return 20

    def can_parse(self, data: Any, format_hint: Optional[str] = None) -> bool:
        if format_hint == "markdown":
            return True
        if isinstance(data, str):
            # Check for markdown patterns
            return bool(re.search(r'^#{1,6}\s', data, re.MULTILINE) or
                       re.search(r'```', data) or
                       re.search(r'^\*|\d+\.', data, re.MULTILINE))
        return False

    def parse(self, data: Any) -> ParseResult:
        """Parse markdown into nodes"""
        try:
            text = str(data).strip()
            nodes = []

            # Split by headings
            sections = re.split(r'^(#{1,6}\s+.+)$', text, flags=re.MULTILINE)

            current_heading = None
            section_idx = 0

            for i, part in enumerate(sections):
                part = part.strip()
                if not part:
                    continue

                # Check if it's a heading
                heading_match = re.match(r'^(#{1,6})\s+(.+)$', part)
                if heading_match:
                    level = len(heading_match.group(1))
                    title = heading_match.group(2)
                    current_heading = title
                    nodes.append(Node(
                        path=f"heading_{section_idx}",
                        content=title,
                        node_type=NodeType.HEADING,
                        metadata={"format": "markdown", "level": level}
                    ))
                    section_idx += 1
                else:
                    # Content section
                    # Check for code blocks
                    code_blocks = list(re.finditer(r'```(\w*)\n(.*?)```', part, re.DOTALL))

                    if code_blocks:
                        for cb in code_blocks:
                            language = cb.group(1) or "text"
                            code = cb.group(2).strip()
                            nodes.append(Node(
                                path=f"code_block_{len(nodes)}",
                                content=code,
                                node_type=NodeType.CODE_BLOCK,
                                metadata={"format": "markdown", "language": language}
                            ))

                    # Remove code blocks and parse remaining as paragraphs
                    remaining = re.sub(r'```.*?```', '', part, flags=re.DOTALL).strip()
                    if remaining:
                        paragraphs = [p.strip() for p in remaining.split('\n\n') if p.strip()]
                        for para in paragraphs:
                            nodes.append(Node(
                                path=f"paragraph_{len(nodes)}",
                                content=para,
                                node_type=NodeType.PARAGRAPH,
                                metadata={"format": "markdown", "heading": current_heading}
                            ))

            return ParseResult(
                nodes=nodes,
                format_name="markdown",
                success=True,
                metadata={"node_count": len(nodes)}
            )
        except Exception as e:
            return ParseResult(
                nodes=[],
                format_name="markdown",
                success=False,
                error=str(e)
            )

    def reconstruct(self, nodes: List[Node], target_format: Optional[str] = None) -> str:
        """Reconstruct markdown from nodes"""
        if not nodes:
            return ""

        parts = []
        for node in nodes:
            if node.node_type == NodeType.HEADING:
                level = node.metadata.get("level", 1)
                parts.append(f"\n{'#' * level} {node.content}\n")
            elif node.node_type == NodeType.CODE_BLOCK:
                language = node.metadata.get("language", "")
                parts.append(f"\n```{language}\n{node.content}\n```\n")
            elif node.node_type == NodeType.PARAGRAPH:
                parts.append(f"\n{node.content}\n")

        return ''.join(parts).strip()


# ============================================================================
# CSV Parser
# ============================================================================

class CSVNodeParser(BaseNodeParser):
    """Parse CSV data into row nodes"""

    @property
    def format_name(self) -> str:
        return "csv"

    @property
    def priority(self) -> int:
        return 15

    def can_parse(self, data: Any, format_hint: Optional[str] = None) -> bool:
        if format_hint == "csv":
            return True
        if isinstance(data, str):
            # Check for CSV pattern (commas, multiple lines)
            lines = data.strip().split('\n')
            if len(lines) > 1:
                return ',' in data
        return False

    def parse(self, data: Any) -> ParseResult:
        """Parse CSV into row nodes"""
        try:
            import csv
            from io import StringIO

            text = str(data).strip()
            reader = csv.DictReader(StringIO(text))

            nodes = []
            for i, row in enumerate(reader):
                nodes.append(Node(
                    path=f"row_{i}",
                    content=dict(row),
                    node_type=NodeType.ROW,
                    metadata={"format": "csv"}
                ))

            return ParseResult(
                nodes=nodes,
                format_name="csv",
                success=True,
                metadata={"node_count": len(nodes)}
            )
        except Exception as e:
            return ParseResult(
                nodes=[],
                format_name="csv",
                success=False,
                error=str(e)
            )

    def reconstruct(self, nodes: List[Node], target_format: Optional[str] = None) -> str:
        """Reconstruct CSV from row nodes"""
        if not nodes:
            return ""

        import csv
        from io import StringIO

        output = StringIO()

        # Get all unique keys for headers
        all_keys = set()
        for node in nodes:
            if isinstance(node.content, dict):
                all_keys.update(node.content.keys())

        headers = sorted(all_keys)
        writer = csv.DictWriter(output, fieldnames=headers)
        writer.writeheader()

        for node in nodes:
            if isinstance(node.content, dict):
                writer.writerow(node.content)

        return output.getvalue()


# ============================================================================
# PDF Parser
# ============================================================================

class PDFNodeParser(BaseNodeParser):
    """Parse PDF documents into nodes by extracting text per page/section"""

    @property
    def format_name(self) -> str:
        return "pdf"

    @property
    def priority(self) -> int:
        return 5  # High priority when format hint matches

    def can_parse(self, data: Any, format_hint: Optional[str] = None) -> bool:
        if format_hint == "pdf":
            return True
        # Check for PDF magic bytes
        if isinstance(data, (bytes, bytearray)):
            return data[:5] == b'%PDF-'
        return False

    def parse(self, data: Any) -> ParseResult:
        """Parse PDF bytes into nodes"""
        try:
            import fitz  # pymupdf

            if isinstance(data, (bytes, bytearray)):
                doc = fitz.open(stream=data, filetype="pdf")
            else:
                return ParseResult(
                    nodes=[],
                    format_name="pdf",
                    success=False,
                    error="PDF parser requires bytes input"
                )

            nodes = []
            total_pages = len(doc)

            for page_num in range(total_pages):
                page = doc[page_num]
                text = page.get_text("text").strip()

                if not text:
                    continue

                # Split page text into paragraphs
                paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

                if len(paragraphs) <= 1:
                    # Single block of text for this page — one section node
                    nodes.append(Node(
                        path=f"page_{page_num + 1}",
                        content=text,
                        node_type=NodeType.SECTION,
                        metadata={
                            "format": "pdf",
                            "page": page_num + 1,
                            "total_pages": total_pages,
                        }
                    ))
                else:
                    # Multiple paragraphs — create individual nodes
                    for i, para in enumerate(paragraphs):
                        nodes.append(Node(
                            path=f"page_{page_num + 1}.paragraph_{i}",
                            content=para,
                            node_type=NodeType.PARAGRAPH,
                            metadata={
                                "format": "pdf",
                                "page": page_num + 1,
                                "total_pages": total_pages,
                            }
                        ))

            doc.close()

            return ParseResult(
                nodes=nodes,
                format_name="pdf",
                success=True,
                metadata={
                    "node_count": len(nodes),
                    "total_pages": total_pages,
                }
            )
        except ImportError:
            return ParseResult(
                nodes=[],
                format_name="pdf",
                success=False,
                error="pymupdf is required for PDF parsing: pip install pymupdf"
            )
        except Exception as e:
            return ParseResult(
                nodes=[],
                format_name="pdf",
                success=False,
                error=str(e)
            )

    def reconstruct(self, nodes: List[Node], target_format: Optional[str] = None) -> str:
        """Reconstruct PDF content as plain text (cannot recreate binary PDF)"""
        if not nodes:
            return ""

        pages: dict[int, list[str]] = {}
        for node in nodes:
            page = node.metadata.get("page", 1)
            if page not in pages:
                pages[page] = []
            pages[page].append(str(node.content))

        parts = []
        for page_num in sorted(pages.keys()):
            parts.append(f"--- Page {page_num} ---")
            parts.append("\n\n".join(pages[page_num]))

        return "\n\n".join(parts)


# ============================================================================
# DOCX Parser
# ============================================================================

class DOCXNodeParser(BaseNodeParser):
    """Parse DOCX documents into nodes using python-docx"""

    @property
    def format_name(self) -> str:
        return "docx"

    @property
    def priority(self) -> int:
        return 5  # High priority when format hint matches

    def can_parse(self, data: Any, format_hint: Optional[str] = None) -> bool:
        if format_hint == "docx":
            return True
        # Check for ZIP/DOCX magic bytes (PK header)
        if isinstance(data, (bytes, bytearray)):
            return data[:2] == b'PK' and b'word/' in data[:2000]
        return False

    def parse(self, data: Any) -> ParseResult:
        """Parse DOCX bytes into nodes"""
        try:
            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError

            if isinstance(data, (bytes, bytearray)):
                doc = Document(io.BytesIO(data))
            else:
                return ParseResult(
                    nodes=[],
                    format_name="docx",
                    success=False,
                    error="DOCX parser requires bytes input"
                )

            nodes = []
            current_heading = None
            para_idx = 0

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue

                style_name = paragraph.style.name if paragraph.style else ""

                if style_name.startswith("Heading"):
                    # Extract heading level from style name (e.g., "Heading 1" -> 1)
                    try:
                        level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1
                    current_heading = text
                    nodes.append(Node(
                        path=f"heading_{len(nodes)}",
                        content=text,
                        node_type=NodeType.HEADING,
                        metadata={
                            "format": "docx",
                            "level": level,
                            "style": style_name,
                        }
                    ))
                elif style_name.startswith("List"):
                    nodes.append(Node(
                        path=f"list_item_{len(nodes)}",
                        content=text,
                        node_type=NodeType.LIST_ITEM,
                        metadata={
                            "format": "docx",
                            "heading": current_heading,
                            "style": style_name,
                        }
                    ))
                else:
                    nodes.append(Node(
                        path=f"paragraph_{para_idx}",
                        content=text,
                        node_type=NodeType.PARAGRAPH,
                        metadata={
                            "format": "docx",
                            "heading": current_heading,
                            "style": style_name,
                        }
                    ))
                    para_idx += 1

            # Extract tables
            for t_idx, table in enumerate(doc.tables):
                # Get headers from first row
                headers = []
                if table.rows:
                    headers = [cell.text.strip() for cell in table.rows[0].cells]

                for r_idx, row in enumerate(table.rows[1:], start=1):
                    row_data = {}
                    for c_idx, cell in enumerate(row.cells):
                        key = headers[c_idx] if c_idx < len(headers) else f"col_{c_idx}"
                        row_data[key] = cell.text.strip()
                    if any(v for v in row_data.values()):
                        nodes.append(Node(
                            path=f"table_{t_idx}.row_{r_idx}",
                            content=row_data,
                            node_type=NodeType.ROW,
                            metadata={
                                "format": "docx",
                                "table_index": t_idx,
                            }
                        ))

            return ParseResult(
                nodes=nodes,
                format_name="docx",
                success=True,
                metadata={"node_count": len(nodes)}
            )
        except ImportError:
            return ParseResult(
                nodes=[],
                format_name="docx",
                success=False,
                error="python-docx is required for DOCX parsing: pip install python-docx"
            )
        except Exception as e:
            return ParseResult(
                nodes=[],
                format_name="docx",
                success=False,
                error=str(e)
            )

    def reconstruct(self, nodes: List[Node], target_format: Optional[str] = None) -> str:
        """Reconstruct DOCX content as plain text (cannot recreate binary DOCX)"""
        if not nodes:
            return ""

        parts = []
        for node in nodes:
            if node.node_type == NodeType.HEADING:
                level = node.metadata.get("level", 1)
                parts.append(f"\n{'#' * level} {node.content}\n")
            elif node.node_type == NodeType.LIST_ITEM:
                parts.append(f"  - {node.content}")
            elif node.node_type == NodeType.ROW:
                if isinstance(node.content, dict):
                    row_str = " | ".join(f"{k}: {v}" for k, v in node.content.items())
                    parts.append(f"  [{row_str}]")
                else:
                    parts.append(f"  {node.content}")
            else:
                parts.append(str(node.content))

        return "\n\n".join(parts)
